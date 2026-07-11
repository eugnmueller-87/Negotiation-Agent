"""Contract file → text — turn an uploaded PDF/DOCX/plaintext into the words to negotiate on.

The demo used to reject binary formats and ask the human to paste the text; this closes
that gap by extracting the text layer server-side, then feeding the SAME ``/intel`` path.

Text-layer extraction only, deliberately:
  - A digitally-generated contract PDF (the common case) carries a text layer — pypdf
    reads it directly, no OCR, no system binary, deploys clean on Nixpacks.
  - A SCANNED PDF is images with no text layer. pypdf returns (almost) nothing. We do
    NOT silently hand back an empty string — we RAISE :class:`NoTextLayer` so the caller
    can tell the human "this looks scanned, paste the text", per the no-silent-failure rule.

The file bytes are untrusted input. We cap the extracted text length (same bound as the
regex extractor) and never execute anything from the document. pypdf/python-docx parse
the container; they do not evaluate embedded scripts.
"""

from __future__ import annotations

import io

from pydantic import BaseModel

# Below this many non-whitespace characters, a PDF is treated as having no usable text
# layer (scanned/image-only) rather than a real-but-short contract. A one-line contract
# is not a real contract; a scanned page yields near-zero extractable characters.
_MIN_TEXT_CHARS = 20

# Same 2 MB bound the regex extractor applies — cap here too so a pathological document
# can't balloon memory. When the extracted text exceeds this we cut it AND set truncated=True
# on the result, so the caller can warn the human (the downstream intake re-check can't fire
# on already-cut text — both caps are identical — so truncation MUST be reported here).
_MAX_TEXT_CHARS = 2_000_000


class ExtractResult(BaseModel):
    """Extracted text plus whether it was cut at the cap. ``truncated`` is surfaced to the
    human — a silent cut on the upload path would violate the no-silent-truncation rule."""

    model_config = {"frozen": True}

    text: str
    truncated: bool = False


class ExtractError(Exception):
    """Base for file-extraction failures — carries a human-safe message and a code."""

    code = "extract_failed"


class UnsupportedFileType(ExtractError):
    """The file extension/content isn't a format we extract (not PDF/DOCX/text)."""

    code = "unsupported_file_type"


class NoTextLayer(ExtractError):
    """A PDF with no extractable text — almost certainly scanned/image-only."""

    code = "no_text_layer"


class CorruptFile(ExtractError):
    """The file couldn't be parsed as its claimed type (truncated, wrong format, corrupt)."""

    code = "corrupt_file"


# A digitally-generated contract is well under this; a file declaring far more pages is a
# complexity/decompression bomb (an attacker packs millions of pages under the 20 MB byte
# cap). Cap the pages we iterate so one crafted PDF can't pin the process for minutes.
_MAX_PDF_PAGES = 500


def _cap(text: str) -> ExtractResult:
    """Cut ``text`` to the char cap, flagging whether a cut happened (surfaced to the human)."""
    if len(text) > _MAX_TEXT_CHARS:
        return ExtractResult(text=text[:_MAX_TEXT_CHARS], truncated=True)
    return ExtractResult(text=text, truncated=False)


def extract_pdf(data: bytes) -> ExtractResult:
    """Extract the text layer from a PDF. Raises :class:`NoTextLayer` if there is none.

    Imports pypdf lazily so the core package keeps its single runtime dependency; the
    import only runs when the ``[web]`` extra is installed and a PDF is actually uploaded.
    Bounded work: at most ``_MAX_PDF_PAGES`` pages, and extraction stops once the accumulated
    text passes the char cap — a crafted many-page PDF can't run for minutes.
    """
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover - exercised only without the extra
        raise ExtractError("PDF support needs the 'web' extra (pypdf)") from e

    try:
        reader = PdfReader(io.BytesIO(data))
        if len(reader.pages) > _MAX_PDF_PAGES:
            raise CorruptFile(
                f"the PDF has more than {_MAX_PDF_PAGES} pages — split it or paste the key clauses."
            )
        parts: list[str] = []
        size = 0
        for page in reader.pages:
            piece = page.extract_text() or ""
            parts.append(piece)
            size += len(piece)
            if size > _MAX_TEXT_CHARS:
                break  # enough text; don't grind through the rest (bounds crafted input)
    except CorruptFile:
        raise
    except Exception as e:  # noqa: BLE001 - pypdf raises varied/undocumented errors on bad input
        raise CorruptFile("the PDF could not be read — it may be corrupt or encrypted") from e

    text = "\n\n".join(p.strip() for p in parts if p.strip())
    if len(text.replace(" ", "").replace("\n", "")) < _MIN_TEXT_CHARS:
        raise NoTextLayer(
            "no text layer found — this looks like a scanned image. "
            "Paste the text, or upload a text-based PDF."
        )
    return _cap(text)


def extract_docx(data: bytes) -> ExtractResult:
    """Extract paragraph + table text from a Word .docx. Raises on a non-docx/corrupt file."""
    try:
        import docx  # python-docx
    except ImportError as e:  # pragma: no cover - exercised only without the extra
        raise ExtractError("DOCX support needs the 'web' extra (python-docx)") from e

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001 - python-docx raises varied errors on a bad file
        raise CorruptFile(
            "the .docx could not be read — it may be corrupt or not a Word file"
        ) from e

    lines = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    text = "\n".join(lines).strip()
    if len(text.replace(" ", "").replace("\n", "")) < _MIN_TEXT_CHARS:
        raise NoTextLayer("the document has no readable text — paste the contract text instead.")
    return _cap(text)


def extract_file(filename: str, data: bytes) -> ExtractResult:
    """Route a file to the right extractor by its extension. The single public entry point.

    Plain-text formats are decoded directly (utf-8, latin-1 fallback). PDF and DOCX go
    to their extractors. Anything else raises :class:`UnsupportedFileType`. Every failure
    is a typed :class:`ExtractError` with a human-safe message — the API maps it to a 4xx.
    Returns an :class:`ExtractResult` whose ``truncated`` flag the caller must surface.
    """
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_pdf(data)
    if name.endswith((".docx",)):
        return extract_docx(data)
    if name.endswith((".txt", ".md", ".markdown", ".csv", ".text")) or "." not in name:
        # decode as text; latin-1 never fails, so it's the last-resort fallback
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1")
        if "\x00" in text[:4096]:
            raise CorruptFile("this file looks binary, not text — paste the contract text instead.")
        return _cap(text)
    raise UnsupportedFileType(
        "unsupported file type — upload a PDF, Word .docx, or plain-text file."
    )
